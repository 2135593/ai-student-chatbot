import os
import re
import uuid
import json
from datetime import datetime, timedelta
from typing import Any, Dict, List, Optional, Tuple

import faiss
import numpy as np
import requests
from docx import Document as DocxDocument
from fastapi import Depends, FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from jose import JWTError, jwt
from langdetect import detect
from passlib.context import CryptContext
from pydantic import BaseModel
from pypdf import PdfReader
from sentence_transformers import SentenceTransformer
from sqlalchemy import Boolean, Column, DateTime, ForeignKey, Integer, String, Text, create_engine
from sqlalchemy.orm import Session, declarative_base, relationship, sessionmaker

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data")
UPLOAD_DIR = os.path.join(DATA_DIR, "uploads")
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)
DB_PATH = os.path.join(DATA_DIR, "chatbot.db")
FAISS_INDEX_PATH = os.path.join(DATA_DIR, "faiss.index")
FAISS_META_PATH = os.path.join(DATA_DIR, "faiss_meta.npy")

SECRET_KEY = os.getenv("SECRET_KEY", "change-this-in-production")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = int(os.getenv("ACCESS_TOKEN_EXPIRE_MINUTES", "720"))
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
FAST_MODEL = os.getenv("FAST_MODEL", "phi4-mini")
QUALITY_MODEL = os.getenv("QUALITY_MODEL", "llama3.1:8b")
EMBED_MODEL = os.getenv("EMBED_MODEL", "all-MiniLM-L6-v2")
TOP_K = int(os.getenv("TOP_K", "6"))
MAX_HISTORY = int(os.getenv("MAX_HISTORY", "16"))

engine = create_engine(f"sqlite:///{DB_PATH}", connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)
Base = declarative_base()
pwd_context = CryptContext(schemes=["pbkdf2_sha256"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/auth/login")


class _DeterministicTestEmbeddingModel:
    """Small deterministic embedding model used only when CHATBOT_TEST_MODE=1."""

    def __init__(self, dimensions: int = 384):
        self.dimensions = dimensions

    def encode(self, texts, convert_to_numpy=True, show_progress_bar=False):
        if isinstance(texts, str):
            texts = [texts]
        vectors = []
        for text in texts:
            vector = np.zeros(self.dimensions, dtype="float32")
            for token in re.findall(r"[a-z0-9]+", text.lower()):
                vector[hash(token) % self.dimensions] += 1.0
            if not vector.any():
                vector[0] = 1.0
            vectors.append(vector)
        return np.vstack(vectors)


if os.getenv("CHATBOT_TEST_MODE") == "1":
    embedding_model = _DeterministicTestEmbeddingModel()
else:
    embedding_model = SentenceTransformer(EMBED_MODEL)

VECTOR_INDEX: Optional[faiss.IndexFlatIP] = None
VECTOR_META: List[int] = []


class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True)
    username = Column(String(100), unique=True, nullable=False, index=True)
    full_name = Column(String(200), nullable=True)
    hashed_password = Column(String(255), nullable=False)
    role = Column(String(20), default="student")
    created_at = Column(DateTime, default=datetime.utcnow)

    conversations = relationship("Conversation", back_populates="user")
    reminders = relationship("Reminder", back_populates="user")


class Conversation(Base):
    __tablename__ = "conversations"
    id = Column(Integer, primary_key=True)
    title = Column(String(255), default="New Chat")
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)
    updated_at = Column(DateTime, default=datetime.utcnow)

    user = relationship("User", back_populates="conversations")
    messages = relationship("Message", back_populates="conversation", cascade="all, delete-orphan")


class Message(Base):
    __tablename__ = "messages"
    id = Column(Integer, primary_key=True)
    conversation_id = Column(Integer, ForeignKey("conversations.id"), nullable=False)
    role = Column(String(20), nullable=False)
    content = Column(Text, nullable=False)
    sources_json = Column(Text, nullable=True)
    created_at = Column(DateTime, default=datetime.utcnow)

    conversation = relationship("Conversation", back_populates="messages")


class SourceDocument(Base):
    __tablename__ = "documents"
    id = Column(Integer, primary_key=True)
    name = Column(String(255), nullable=False)
    file_type = Column(String(50), nullable=False)
    file_path = Column(String(500), nullable=True)
    original_filename = Column(String(255), nullable=True)
    uploaded_by = Column(Integer, ForeignKey("users.id"), nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)

    chunks = relationship("DocumentChunk", back_populates="document", cascade="all, delete-orphan")


class DocumentChunk(Base):
    __tablename__ = "document_chunks"
    id = Column(Integer, primary_key=True)
    document_id = Column(Integer, ForeignKey("documents.id"), nullable=False)
    chunk_index = Column(Integer, nullable=False)
    text = Column(Text, nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)

    document = relationship("SourceDocument", back_populates="chunks")


class FAQ(Base):
    __tablename__ = "faqs"
    id = Column(Integer, primary_key=True)
    question = Column(Text, nullable=False)
    answer = Column(Text, nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)
    updated_at = Column(DateTime, default=datetime.utcnow)


class Reminder(Base):
    __tablename__ = "reminders"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    title = Column(String(255), nullable=False)
    due_date = Column(String(20), nullable=True)
    notes = Column(Text, nullable=True)
    is_completed = Column(Boolean, default=False)
    created_at = Column(DateTime, default=datetime.utcnow)
    updated_at = Column(DateTime, default=datetime.utcnow)

    user = relationship("User", back_populates="reminders")


Base.metadata.create_all(bind=engine)


def ensure_schema():
    with engine.begin() as conn:
        rows = conn.exec_driver_sql("PRAGMA table_info(messages)").fetchall()
        message_cols = {row[1] for row in rows}
        if "sources_json" not in message_cols:
            conn.exec_driver_sql("ALTER TABLE messages ADD COLUMN sources_json TEXT")
        rows = conn.exec_driver_sql("PRAGMA table_info(documents)").fetchall()
        document_cols = {row[1] for row in rows}
        if "file_path" not in document_cols:
            conn.exec_driver_sql("ALTER TABLE documents ADD COLUMN file_path VARCHAR(500)")
        if "original_filename" not in document_cols:
            conn.exec_driver_sql("ALTER TABLE documents ADD COLUMN original_filename VARCHAR(255)")


ensure_schema()


class TokenResponse(BaseModel):
    access_token: str
    token_type: str
    user: Dict[str, Any]


class UserOut(BaseModel):
    id: int
    username: str
    full_name: Optional[str]
    role: str

    class Config:
        from_attributes = True


class ConversationOut(BaseModel):
    id: int
    title: str
    created_at: datetime
    updated_at: datetime

    class Config:
        from_attributes = True


class MessageOut(BaseModel):
    id: int
    role: str
    content: str
    created_at: datetime
    sources: List[Dict[str, Any]] = []

    class Config:
        from_attributes = True


class ChatRequest(BaseModel):
    conversation_id: Optional[int] = None
    message: str
    mode: str = "quality"


class SourceOut(BaseModel):
    kind: str
    item_id: int
    title: str
    snippet: str
    score: float
    citation_id: str
    chunk_index: Optional[int] = None
    confidence: str = "medium"


class ChatResponse(BaseModel):
    conversation_id: int
    answer: str
    sources: List[SourceOut]


class DocumentOut(BaseModel):
    id: int
    name: str
    file_type: str
    created_at: datetime
    chunk_count: int
    file_available: bool = False


class FAQIn(BaseModel):
    question: str
    answer: str


class FAQOut(BaseModel):
    id: int
    question: str
    answer: str
    created_at: datetime
    updated_at: datetime

    class Config:
        from_attributes = True


class RegisterIn(BaseModel):
    username: str
    password: str
    full_name: Optional[str] = None


class UserCreateIn(BaseModel):
    username: str
    password: str
    full_name: Optional[str] = None
    role: str = "student"


class UserUpdateIn(BaseModel):
    full_name: Optional[str] = None
    role: Optional[str] = None
    password: Optional[str] = None


class DocumentUpdateIn(BaseModel):
    name: str


class ReminderIn(BaseModel):
    title: str
    due_date: Optional[str] = None
    notes: Optional[str] = None
    is_completed: bool = False


class ReminderOut(BaseModel):
    id: int
    title: str
    due_date: Optional[str]
    notes: Optional[str]
    is_completed: bool
    created_at: datetime
    updated_at: datetime

    class Config:
        from_attributes = True


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


def hash_password(password: str) -> str:
    return pwd_context.hash(password)


def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)


def validate_username(username: str) -> str:
    username = clean_text(username).lower()
    if not re.fullmatch(r"[a-z0-9_.-]{3,40}", username):
        raise HTTPException(status_code=400, detail="Username must be 3-40 characters using letters, numbers, dots, underscores or hyphens")
    return username


def validate_password(password: str) -> str:
    if len(password or "") < 8:
        raise HTTPException(status_code=400, detail="Password must be at least 8 characters")
    return password


def validate_role(role: str) -> str:
    role = (role or "student").strip().lower()
    if role not in {"student", "admin"}:
        raise HTTPException(status_code=400, detail="Role must be either student or admin")
    return role


def create_access_token(data: dict, expires_delta: Optional[timedelta] = None) -> str:
    to_encode = data.copy()
    expire = datetime.utcnow() + (expires_delta or timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES))
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)


def get_current_user(token: str = Depends(oauth2_scheme), db: Session = Depends(get_db)) -> User:
    credentials_exception = HTTPException(status_code=401, detail="Could not validate credentials")
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username = payload.get("sub")
        if not username:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
    user = db.query(User).filter(User.username == username).first()
    if not user:
        raise credentials_exception
    return user


def require_admin(current_user: User = Depends(get_current_user)) -> User:
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    return current_user


def normalise_embeddings(vectors: np.ndarray) -> np.ndarray:
    vectors = vectors.astype("float32")
    norms = np.linalg.norm(vectors, axis=1, keepdims=True)
    norms[norms == 0] = 1e-12
    return vectors / norms


def save_vector_state() -> None:
    global VECTOR_INDEX, VECTOR_META
    if VECTOR_INDEX is None:
        return
    faiss.write_index(VECTOR_INDEX, FAISS_INDEX_PATH)
    np.save(FAISS_META_PATH, np.array(VECTOR_META, dtype=np.int64))


def load_vector_state() -> None:
    global VECTOR_INDEX, VECTOR_META
    if os.path.exists(FAISS_INDEX_PATH) and os.path.exists(FAISS_META_PATH):
        VECTOR_INDEX = faiss.read_index(FAISS_INDEX_PATH)
        VECTOR_META = np.load(FAISS_META_PATH).tolist()
    else:
        VECTOR_INDEX = None
        VECTOR_META = []


def rebuild_vector_index(db: Session) -> None:
    global VECTOR_INDEX, VECTOR_META
    chunks = db.query(DocumentChunk).order_by(DocumentChunk.id.asc()).all()
    if not chunks:
        VECTOR_INDEX = None
        VECTOR_META = []
        if os.path.exists(FAISS_INDEX_PATH):
            os.remove(FAISS_INDEX_PATH)
        if os.path.exists(FAISS_META_PATH):
            os.remove(FAISS_META_PATH)
        return
    texts = [c.text for c in chunks]
    embeddings = embedding_model.encode(texts, convert_to_numpy=True, show_progress_bar=False)
    embeddings = normalise_embeddings(np.array(embeddings))
    index = faiss.IndexFlatIP(embeddings.shape[1])
    index.add(embeddings)
    VECTOR_INDEX = index
    VECTOR_META = [c.id for c in chunks]
    save_vector_state()


def ensure_seed_users(db: Session) -> None:
    defaults = [
        ("admin", "Admin User", "admin123", "admin"),
        ("student", "Student User", "student123", "student"),
    ]
    for username, full_name, password, role in defaults:
        user = db.query(User).filter(User.username == username).first()
        if not user:
            db.add(User(username=username, full_name=full_name, hashed_password=hash_password(password), role=role))
            continue
        if not user.hashed_password or not (user.hashed_password.startswith("$pbkdf2-sha256$") or user.hashed_password.startswith("$pbkdf2_sha256$")):
            user.hashed_password = hash_password(password)
            user.full_name = full_name
            user.role = role
    db.commit()


def ensure_seed_faqs(db: Session) -> None:
    if db.query(FAQ).count() > 0:
        return
    db.add_all(
        [
            FAQ(question="How do I use the chatbot?", answer="Log in, open a chat, and ask a question. Admins can upload documents and manage FAQs."),
            FAQ(question="Who can upload documents?", answer="Only admin users can upload and delete knowledge-base documents."),
            FAQ(question="Can I create reminders?", answer="Yes. Both students and admins can add, edit, complete, and delete their own reminders."),
        ]
    )
    db.commit()


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def read_upload_text(file: UploadFile, raw: bytes) -> str:
    filename = (file.filename or "uploaded_file").lower()
    if filename.endswith(".txt"):
        return raw.decode("utf-8", errors="ignore")
    if filename.endswith(".pdf"):
        tmp = os.path.join(DATA_DIR, f"tmp_{uuid.uuid4().hex}.pdf")
        with open(tmp, "wb") as f:
            f.write(raw)
        try:
            reader = PdfReader(tmp)
            return "\n\n".join((page.extract_text() or "") for page in reader.pages)
        finally:
            if os.path.exists(tmp):
                os.remove(tmp)
    if filename.endswith(".docx"):
        tmp = os.path.join(DATA_DIR, f"tmp_{uuid.uuid4().hex}.docx")
        with open(tmp, "wb") as f:
            f.write(raw)
        try:
            doc = DocxDocument(tmp)
            return "\n".join(p.text for p in doc.paragraphs)
        finally:
            if os.path.exists(tmp):
                os.remove(tmp)
    raise HTTPException(status_code=400, detail="Unsupported file type. Use .txt, .pdf, or .docx")


def chunk_text(text: str, target_words: int = 180, overlap_words: int = 40) -> List[str]:
    text = clean_text(text)
    if not text:
        return []
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: List[str] = []
    current_words: List[str] = []
    for para in paragraphs:
        para_words = para.split()
        if len(current_words) + len(para_words) <= target_words:
            current_words.extend(para_words)
        else:
            if current_words:
                chunks.append(" ".join(current_words).strip())
                overlap = current_words[-overlap_words:] if len(current_words) > overlap_words else current_words[:]
                current_words = overlap + para_words
            else:
                for start in range(0, len(para_words), max(target_words - overlap_words, 1)):
                    piece = para_words[start : start + target_words]
                    if piece:
                        chunks.append(" ".join(piece).strip())
                current_words = []
    if current_words:
        chunks.append(" ".join(current_words).strip())

    unique_chunks = []
    seen = set()
    for chunk in chunks:
        key = chunk.lower()
        if key not in seen and len(chunk.split()) >= 20:
            seen.add(key)
            unique_chunks.append(chunk)
    return unique_chunks


def extract_snippet(text: str, query: str, max_chars: int = 240) -> str:
    clean = clean_text(text)
    query_terms = [t for t in re.findall(r"\w+", query.lower()) if len(t) > 2]
    lower = clean.lower()
    for term in query_terms:
        idx = lower.find(term)
        if idx != -1:
            start = max(0, idx - 70)
            end = min(len(clean), idx + 170)
            snippet = clean[start:end]
            if start > 0:
                snippet = "..." + snippet
            if end < len(clean):
                snippet += "..."
            return snippet[:max_chars]
    return clean[:max_chars] + ("..." if len(clean) > max_chars else "")


def keyword_overlap_score(query: str, text: str) -> float:
    q_terms = {t for t in re.findall(r"\w+", query.lower()) if len(t) > 2}
    t_terms = set(re.findall(r"\w+", text.lower()))
    if not q_terms:
        return 0.0
    return len(q_terms & t_terms) / max(len(q_terms), 1)


def retrieve_document_chunks(db: Session, query: str, k: int = TOP_K) -> List[Tuple[DocumentChunk, SourceDocument, float]]:
    global VECTOR_INDEX, VECTOR_META
    if VECTOR_INDEX is None or not VECTOR_META:
        return []
    query_embedding = embedding_model.encode([query], convert_to_numpy=True, show_progress_bar=False)
    query_embedding = normalise_embeddings(np.array(query_embedding))
    search_k = min(max(k * 3, k), len(VECTOR_META))
    scores, indices = VECTOR_INDEX.search(query_embedding, search_k)
    ranked: List[Tuple[DocumentChunk, SourceDocument, float]] = []
    for sim_score, idx in zip(scores[0].tolist(), indices[0].tolist()):
        if idx < 0 or idx >= len(VECTOR_META):
            continue
        chunk_id = VECTOR_META[idx]
        chunk = db.query(DocumentChunk).filter(DocumentChunk.id == chunk_id).first()
        if not chunk:
            continue
        doc = db.query(SourceDocument).filter(SourceDocument.id == chunk.document_id).first()
        if not doc:
            continue
        hybrid = float(sim_score) + 0.35 * keyword_overlap_score(query, chunk.text)
        ranked.append((chunk, doc, hybrid))
    ranked.sort(key=lambda item: item[2], reverse=True)
    unique: List[Tuple[DocumentChunk, SourceDocument, float]] = []
    seen = set()
    for item in ranked:
        if item[0].id in seen:
            continue
        seen.add(item[0].id)
        unique.append(item)
        if len(unique) >= k:
            break
    return unique


def retrieve_faqs(db: Session, query: str, k: int = 3) -> List[Tuple[FAQ, float]]:
    faqs = db.query(FAQ).order_by(FAQ.updated_at.desc()).all()
    scored: List[Tuple[FAQ, float]] = []
    for faq in faqs:
        score = keyword_overlap_score(query, f"{faq.question} {faq.answer}")
        if score > 0:
            scored.append((faq, score))
    scored.sort(key=lambda item: item[1], reverse=True)
    return scored[:k]


def safe_detect_language(text: str) -> str:
    try:
        return detect(text)
    except Exception:
        return "en"


def select_model(mode: str) -> str:
    return FAST_MODEL if mode.lower() == "fast" else QUALITY_MODEL


def build_chat_title(message: str) -> str:
    message = clean_text(message).replace("\n", " ")
    return message[:48] + ("..." if len(message) > 48 else "") if message else "New Chat"


def call_ollama(model: str, prompt: str) -> str:
    try:
        response = requests.post(
            f"{OLLAMA_BASE_URL.rstrip('/')}/api/generate",
            json={"model": model, "prompt": prompt, "stream": False, "options": {"temperature": 0.2}},
            timeout=180,
        )
        response.raise_for_status()
        return (response.json().get("response") or "").strip()
    except requests.RequestException as exc:
        raise HTTPException(status_code=503, detail=f"Ollama is not running or the selected model is unavailable. Start Ollama and pull the configured model. Technical detail: {exc}")


def source_confidence(score: float, kind: str) -> str:
    if kind == "faq":
        return "high" if score >= 0.35 else "medium"
    if score >= 0.72:
        return "high"
    if score >= 0.48:
        return "medium"
    return "low"


def format_sources_for_answer(sources: List[SourceOut]) -> str:
    if not sources:
        return ""
    lines = []
    seen = set()
    for src in sources:
        key = (src.kind, src.item_id, src.citation_id)
        if key in seen:
            continue
        seen.add(key)
        label = f"[{src.citation_id}] {src.title}"
        if src.chunk_index is not None:
            label += f" (chunk {src.chunk_index + 1})"
        lines.append(f"- {label}")
    return "\n".join(lines)


def build_prompt(user_message: str, history: List[Message], doc_sources: List[Tuple[DocumentChunk, SourceDocument, float]], faq_sources: List[Tuple[FAQ, float]], language: str) -> str:
    history_text = "\n".join(
        [f"{'User' if msg.role == 'user' else 'Assistant'}: {msg.content}" for msg in history[-MAX_HISTORY:]]
    ) or "No previous messages."

    doc_context = "\n\n".join(
        [f"[D{i}] Document: {doc.name} | Chunk: {chunk.chunk_index + 1}\n{chunk.text}" for i, (chunk, doc, _) in enumerate(doc_sources, start=1)]
    ) or "No relevant uploaded document content found."

    faq_context = "\n\n".join(
        [f"[F{i}] FAQ: {faq.question}\nAnswer: {faq.answer}" for i, (faq, _) in enumerate(faq_sources, start=1)]
    ) or "No relevant FAQs found."

    return f"""
You are a helpful local student-support assistant.

Rules:
- Reply in the same language as the user. Detected language: {language}
- Prefer uploaded document context when relevant.
- FAQs are trusted fixed answers and can also be used.
- If the answer is not clearly supported, say the available material does not confirm it.
- Keep the answer practical and concise.
- Use only the supplied Document context and FAQ context for factual academic/course information.
- Cite claims using the bracket IDs exactly as provided, for example [D1] or [F1].
- If you used document or FAQ material, end with a short Sources section listing those IDs.

Conversation history:
{history_text}

Document context:
{doc_context}

FAQ context:
{faq_context}

User question:
{user_message}
""".strip()


def check_ollama_status() -> Dict[str, Any]:
    try:
        response = requests.get(f"{OLLAMA_BASE_URL.rstrip('/')}/api/tags", timeout=3)
        response.raise_for_status()
        models = [item.get("name") for item in response.json().get("models", [])]
        return {"available": True, "models": models}
    except requests.RequestException as exc:
        return {"available": False, "error": str(exc), "models": []}


app = FastAPI(title="Local AI Chatbot", version="3.1.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
def startup_event() -> None:
    load_vector_state()
    db = SessionLocal()
    try:
        ensure_seed_users(db)
        ensure_seed_faqs(db)
        if db.query(DocumentChunk).count() > 0 and (VECTOR_INDEX is None or len(VECTOR_META) != db.query(DocumentChunk).count()):
            rebuild_vector_index(db)
    finally:
        db.close()


@app.post("/auth/login", response_model=TokenResponse)
def login(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(get_db)):
    user = db.query(User).filter(User.username == form_data.username).first()
    if not user or not verify_password(form_data.password, user.hashed_password):
        raise HTTPException(status_code=401, detail="Invalid username or password")
    token = create_access_token({"sub": user.username, "role": user.role})
    return {
        "access_token": token,
        "token_type": "bearer",
        "user": {"id": user.id, "username": user.username, "full_name": user.full_name, "role": user.role},
    }


@app.post("/auth/register", response_model=UserOut)
def register(payload: RegisterIn, db: Session = Depends(get_db)):
    username = validate_username(payload.username)
    validate_password(payload.password)
    if db.query(User).filter(User.username == username).first():
        raise HTTPException(status_code=409, detail="Username already exists")
    user = User(username=username, full_name=clean_text(payload.full_name or username), hashed_password=hash_password(payload.password), role="student")
    db.add(user)
    db.commit()
    db.refresh(user)
    return user


@app.get("/auth/me", response_model=UserOut)
def auth_me(current_user: User = Depends(get_current_user)):
    return current_user


@app.get("/admin/users", response_model=List[UserOut])
def list_users(current_user: User = Depends(require_admin), db: Session = Depends(get_db)):
    return db.query(User).order_by(User.created_at.desc()).all()


@app.post("/admin/users", response_model=UserOut)
def create_user(payload: UserCreateIn, current_user: User = Depends(require_admin), db: Session = Depends(get_db)):
    username = validate_username(payload.username)
    validate_password(payload.password)
    role = validate_role(payload.role)
    if db.query(User).filter(User.username == username).first():
        raise HTTPException(status_code=409, detail="Username already exists")
    user = User(username=username, full_name=clean_text(payload.full_name or username), hashed_password=hash_password(payload.password), role=role)
    db.add(user)
    db.commit()
    db.refresh(user)
    return user


@app.put("/admin/users/{user_id}", response_model=UserOut)
def update_user(user_id: int, payload: UserUpdateIn, current_user: User = Depends(require_admin), db: Session = Depends(get_db)):
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    if payload.full_name is not None:
        user.full_name = clean_text(payload.full_name)
    if payload.role is not None:
        if user.id == current_user.id and payload.role != "admin":
            raise HTTPException(status_code=400, detail="You cannot remove your own admin role")
        user.role = validate_role(payload.role)
    if payload.password:
        validate_password(payload.password)
        user.hashed_password = hash_password(payload.password)
    db.commit()
    db.refresh(user)
    return user


@app.delete("/admin/users/{user_id}")
def delete_user(user_id: int, current_user: User = Depends(require_admin), db: Session = Depends(get_db)):
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    if user.id == current_user.id:
        raise HTTPException(status_code=400, detail="You cannot delete your own account")
    db.delete(user)
    db.commit()
    return {"message": "User deleted"}


@app.get("/conversations", response_model=List[ConversationOut])
def list_conversations(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    return db.query(Conversation).filter(Conversation.user_id == current_user.id).order_by(Conversation.updated_at.desc()).all()


@app.post("/conversations", response_model=ConversationOut)
def create_conversation(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    convo = Conversation(title="New Chat", user_id=current_user.id)
    db.add(convo)
    db.commit()
    db.refresh(convo)
    return convo


@app.delete("/conversations/{conversation_id}")
def delete_conversation(conversation_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    convo = db.query(Conversation).filter(Conversation.id == conversation_id, Conversation.user_id == current_user.id).first()
    if not convo:
        raise HTTPException(status_code=404, detail="Conversation not found")
    db.delete(convo)
    db.commit()
    return {"message": "Conversation deleted"}


@app.get("/conversations/{conversation_id}/messages", response_model=List[MessageOut])
def get_messages(conversation_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    convo = db.query(Conversation).filter(Conversation.id == conversation_id, Conversation.user_id == current_user.id).first()
    if not convo:
        raise HTTPException(status_code=404, detail="Conversation not found")
    rows = db.query(Message).filter(Message.conversation_id == conversation_id).order_by(Message.created_at.asc()).all()
    output = []
    for row in rows:
        try:
            sources = json.loads(row.sources_json or "[]")
        except json.JSONDecodeError:
            sources = []
        output.append(MessageOut(id=row.id, role=row.role, content=row.content, created_at=row.created_at, sources=sources))
    return output


@app.post("/chat", response_model=ChatResponse)
def chat(payload: ChatRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    text = clean_text(payload.message)
    if not text:
        raise HTTPException(status_code=400, detail="Message cannot be empty")

    convo: Optional[Conversation] = None
    if payload.conversation_id:
        convo = db.query(Conversation).filter(Conversation.id == payload.conversation_id, Conversation.user_id == current_user.id).first()
        if not convo:
            raise HTTPException(status_code=404, detail="Conversation not found")
    else:
        convo = Conversation(title=build_chat_title(text), user_id=current_user.id)
        db.add(convo)
        db.commit()
        db.refresh(convo)

    if convo.title == "New Chat":
        convo.title = build_chat_title(text)

    db.add(Message(conversation_id=convo.id, role="user", content=text))
    db.commit()

    history = db.query(Message).filter(Message.conversation_id == convo.id).order_by(Message.created_at.asc()).all()
    doc_sources = retrieve_document_chunks(db, text, TOP_K)
    faq_sources = retrieve_faqs(db, text, 3)

    source_outputs: List[SourceOut] = []
    for index, (chunk, doc, score) in enumerate(doc_sources, start=1):
        rounded = round(score, 4)
        source_outputs.append(
            SourceOut(
                kind="document",
                item_id=doc.id,
                title=doc.name,
                snippet=extract_snippet(chunk.text, text),
                score=rounded,
                citation_id=f"D{index}",
                chunk_index=chunk.chunk_index,
                confidence=source_confidence(score, "document"),
            )
        )
    for index, (faq, score) in enumerate(faq_sources, start=1):
        rounded = round(score, 4)
        source_outputs.append(
            SourceOut(
                kind="faq",
                item_id=faq.id,
                title=faq.question,
                snippet=extract_snippet(faq.answer, text),
                score=rounded,
                citation_id=f"F{index}",
                confidence=source_confidence(score, "faq"),
            )
        )

    has_usable_source = any(src.confidence in {"high", "medium"} for src in source_outputs)
    if not has_usable_source:
        answer = (
            "I could not find this answer in the uploaded documents or FAQ content. "
            "To avoid giving unreliable information, please check the official module guide, timetable, policy document, or contact your tutor/student support team."
        )
    else:
        prompt = build_prompt(text, history[:-1], doc_sources, faq_sources, safe_detect_language(text))
        answer = call_ollama(select_model(payload.mode), prompt)
        if "Sources:" not in answer:
            answer += "\n\nSources:\n" + format_sources_for_answer(source_outputs)

    db.add(Message(conversation_id=convo.id, role="assistant", content=answer, sources_json=json.dumps([src.dict() for src in source_outputs])))
    convo.updated_at = datetime.utcnow()
    db.commit()

    return ChatResponse(conversation_id=convo.id, answer=answer, sources=source_outputs)


@app.get("/documents", response_model=List[DocumentOut])
def list_documents(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    docs = db.query(SourceDocument).order_by(SourceDocument.created_at.desc()).all()
    return [DocumentOut(id=doc.id, name=doc.name, file_type=doc.file_type, created_at=doc.created_at, chunk_count=len(doc.chunks), file_available=bool(doc.file_path and os.path.exists(doc.file_path))) for doc in docs]


@app.post("/documents/upload")
def upload_document(file: UploadFile = File(...), display_name: Optional[str] = Form(None), current_user: User = Depends(require_admin), db: Session = Depends(get_db)):
    raw = file.file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")
    text = read_upload_text(file, raw)
    chunks = chunk_text(text)
    if not chunks:
        raise HTTPException(status_code=400, detail="Could not extract enough readable text from the file")

    original_name = file.filename or f"document_{uuid.uuid4().hex[:8]}"
    safe_name = re.sub(r"[^a-zA-Z0-9_.-]", "_", original_name)
    stored_name = f"{uuid.uuid4().hex}_{safe_name}"
    stored_path = os.path.join(UPLOAD_DIR, stored_name)
    with open(stored_path, "wb") as stored_file:
        stored_file.write(raw)
    ext = os.path.splitext(original_name)[1].replace(".", "").lower() or "unknown"
    citation_name = clean_text(display_name or original_name)
    document = SourceDocument(name=citation_name, file_type=ext, file_path=stored_path, original_filename=original_name, uploaded_by=current_user.id)
    db.add(document)
    db.commit()
    db.refresh(document)

    for idx, chunk in enumerate(chunks):
        db.add(DocumentChunk(document_id=document.id, chunk_index=idx, text=chunk))
    db.commit()
    rebuild_vector_index(db)
    return {"message": "Document uploaded successfully", "document_id": document.id, "chunks_created": len(chunks)}


@app.put("/documents/{document_id}", response_model=DocumentOut)
def update_document(document_id: int, payload: DocumentUpdateIn, current_user: User = Depends(require_admin), db: Session = Depends(get_db)):
    doc = db.query(SourceDocument).filter(SourceDocument.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    new_name = clean_text(payload.name)
    if not new_name:
        raise HTTPException(status_code=400, detail="Document name cannot be empty")
    doc.name = new_name
    db.commit()
    db.refresh(doc)
    return DocumentOut(id=doc.id, name=doc.name, file_type=doc.file_type, created_at=doc.created_at, chunk_count=len(doc.chunks), file_available=bool(doc.file_path and os.path.exists(doc.file_path)))


def media_type_for_file(file_type: str) -> str:
    file_type = (file_type or "").lower()
    if file_type == "pdf":
        return "application/pdf"
    if file_type == "docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if file_type == "txt":
        return "text/plain; charset=utf-8"
    return "application/octet-stream"


@app.get("/documents/{document_id}/download")
def download_document(document_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc = db.query(SourceDocument).filter(SourceDocument.id == document_id).first()
    if not doc or not doc.file_path or not os.path.exists(doc.file_path):
        raise HTTPException(status_code=404, detail="Original document file is not available")
    filename = doc.original_filename or doc.name
    return FileResponse(doc.file_path, filename=filename, media_type=media_type_for_file(doc.file_type))


@app.get("/documents/{document_id}/preview")
def preview_document(document_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc = db.query(SourceDocument).filter(SourceDocument.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    chunks = db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).order_by(DocumentChunk.chunk_index.asc()).limit(5).all()
    return {"id": doc.id, "name": doc.name, "file_type": doc.file_type, "text_preview": "\n\n".join(chunk.text for chunk in chunks)}


@app.delete("/documents/{document_id}")
def delete_document(document_id: int, current_user: User = Depends(require_admin), db: Session = Depends(get_db)):
    doc = db.query(SourceDocument).filter(SourceDocument.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if doc.file_path and os.path.exists(doc.file_path):
        try:
            os.remove(doc.file_path)
        except OSError:
            pass
    db.delete(doc)
    db.commit()
    rebuild_vector_index(db)
    return {"message": "Document deleted"}


@app.get("/faqs", response_model=List[FAQOut])
def list_faqs(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    return db.query(FAQ).order_by(FAQ.updated_at.desc()).all()


@app.post("/faqs", response_model=FAQOut)
def create_faq(payload: FAQIn, current_user: User = Depends(require_admin), db: Session = Depends(get_db)):
    question = clean_text(payload.question)
    answer = clean_text(payload.answer)
    if not question or not answer:
        raise HTTPException(status_code=400, detail="FAQ question and answer are required")
    faq = FAQ(question=question, answer=answer, updated_at=datetime.utcnow())
    db.add(faq)
    db.commit()
    db.refresh(faq)
    return faq


@app.put("/faqs/{faq_id}", response_model=FAQOut)
def update_faq(faq_id: int, payload: FAQIn, current_user: User = Depends(require_admin), db: Session = Depends(get_db)):
    faq = db.query(FAQ).filter(FAQ.id == faq_id).first()
    if not faq:
        raise HTTPException(status_code=404, detail="FAQ not found")
    question = clean_text(payload.question)
    answer = clean_text(payload.answer)
    if not question or not answer:
        raise HTTPException(status_code=400, detail="FAQ question and answer are required")
    faq.question = question
    faq.answer = answer
    faq.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(faq)
    return faq


@app.delete("/faqs/{faq_id}")
def delete_faq(faq_id: int, current_user: User = Depends(require_admin), db: Session = Depends(get_db)):
    faq = db.query(FAQ).filter(FAQ.id == faq_id).first()
    if not faq:
        raise HTTPException(status_code=404, detail="FAQ not found")
    db.delete(faq)
    db.commit()
    return {"message": "FAQ deleted"}


@app.get("/reminders", response_model=List[ReminderOut])
def list_reminders(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    return db.query(Reminder).filter(Reminder.user_id == current_user.id).order_by(Reminder.created_at.desc()).all()


@app.post("/reminders", response_model=ReminderOut)
def create_reminder(payload: ReminderIn, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    title = clean_text(payload.title)
    if not title:
        raise HTTPException(status_code=400, detail="Reminder title is required")
    reminder = Reminder(
        user_id=current_user.id,
        title=title,
        due_date=(payload.due_date or "").strip() or None,
        notes=clean_text(payload.notes or "") or None,
        is_completed=payload.is_completed,
        updated_at=datetime.utcnow(),
    )
    db.add(reminder)
    db.commit()
    db.refresh(reminder)
    return reminder


@app.put("/reminders/{reminder_id}", response_model=ReminderOut)
def update_reminder(reminder_id: int, payload: ReminderIn, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    reminder = db.query(Reminder).filter(Reminder.id == reminder_id, Reminder.user_id == current_user.id).first()
    if not reminder:
        raise HTTPException(status_code=404, detail="Reminder not found")
    title = clean_text(payload.title)
    if not title:
        raise HTTPException(status_code=400, detail="Reminder title is required")
    reminder.title = title
    reminder.due_date = (payload.due_date or "").strip() or None
    reminder.notes = clean_text(payload.notes or "") or None
    reminder.is_completed = payload.is_completed
    reminder.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(reminder)
    return reminder


@app.delete("/reminders/{reminder_id}")
def delete_reminder(reminder_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    reminder = db.query(Reminder).filter(Reminder.id == reminder_id, Reminder.user_id == current_user.id).first()
    if not reminder:
        raise HTTPException(status_code=404, detail="Reminder not found")
    db.delete(reminder)
    db.commit()
    return {"message": "Reminder deleted"}


@app.get("/health")
def health(db: Session = Depends(get_db)):
    return {
        "status": "ok",
        "ollama_base_url": OLLAMA_BASE_URL,
        "fast_model": FAST_MODEL,
        "quality_model": QUALITY_MODEL,
        "embed_model": EMBED_MODEL,
        "users": db.query(User).count(),
        "documents": db.query(SourceDocument).count(),
        "faqs": db.query(FAQ).count(),
        "reminders": db.query(Reminder).count(),
        "chunks": db.query(DocumentChunk).count(),
        "vector_index_loaded": VECTOR_INDEX is not None,
        "ollama": check_ollama_status(),
    }
